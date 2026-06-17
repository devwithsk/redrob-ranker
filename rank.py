#!/usr/bin/env python3
"""
Redrob Hackathon -- Candidate Ranking Engine (v2)
==================================================

Single-command usage (per submission_spec.docx Section 10.3):
    python main.py --candidates ./candidates.jsonl.gz --jd ./job_description.md --out ./team_submission.csv

All three flags are optional -- if omitted, the script auto-discovers
candidates.jsonl(.gz) and job_description.md next to this file, and writes
team_submission.csv next to this file.

WHAT CHANGED FROM v1 AND WHY
-----------------------------
1. Reads the real job_description.md instead of scoring everyone against a
   hand-typed buzzword string. Falls back to a curated keyword blob ONLY if
   the file truly can't be found, and prints a loud warning when it does.
2. Auto-detects gzip by magic bytes, not filename -- works whether you hand
   it candidates.jsonl.gz, a pre-unzipped candidates.jsonl, or an
   extension-less file. Removes a real Stage-3 reproducibility risk.
3. "Today" for inactivity math is derived from the freshest last_active_date
   actually present in the dataset (override-able via --reference-date),
   instead of a hardcoded constant that silently drifts if Stage 3
   reproduces this on a different real-world date.
4. Honeypot detection now requires corroborated evidence (skill-duration
   impossibility, OR multiple "expert skill with zero time used" entries)
   instead of a single brittle rule, and reports a rate so you can sanity
   check against the <=10% Stage-3 cutoff yourself. A guarded, schema-aware
   hook for the "company founded after employment started" pattern is
   included but is a documented no-op until you confirm the exact field
   name in candidate_schema.json.
5. Inactivity is now a continuous soft penalty (matches the signals doc's
   own guidance to use behavioral signals "as a multiplier or modifier"),
   not a hard exclude that can drop an otherwise-perfect candidate for being
   181 days inactive instead of 179.
6. Text similarity uses sublinear_tf to blunt keyword-stuffing (the dataset's
   own stated trap category), PLUS an explicit skill-overlap signal so the
   reasoning column can name real matched skills instead of an opaque
   cosine-similarity float. A separate stuffing-suspect penalty flags
   profiles with many shallow/unused skills.
7. Behavioral score now uses 7 of the 23 signals (was 2), including
   open_to_work_flag, which v1 never touched despite it being the most
   direct "are they actually hireable right now" signal in the dataset.
8. Role keyword lists expanded (v1 let an "Operations Manager" reach rank 96
   of an AI/ML ranking task).
9. Reasoning is generated per-candidate from the actual computed
   strengths/concerns for that row -- it varies in length and content, and
   explicitly states concerns for lower-ranked candidates instead of always
   saying "Strong tech score" regardless of the number in parentheses.
10. A lightweight internal self-check mirrors the official validator's core
    rules before writing the CSV, so a logic bug fails loudly here instead
    of silently at Stage 1.

Dependencies: pandas, scikit-learn (stdlib otherwise).
Compute: single-pass per-candidate loop + one TF-IDF fit/transform over the
full pool. No GPU, no network calls -- satisfies Section 3's constraints.
"""

import argparse
import gzip
import json
import math
import sys
import time
from datetime import datetime, date
from pathlib import Path

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

SCRIPT_DIR = Path(__file__).resolve().parent

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Used only if job_description.md cannot be found at all (see load_job_description).
FALLBACK_JD_TEXT = """
modern ML systems embeddings retrieval ranking LLMs fine-tuning
sentence-transformers OpenAI embeddings BGE E5 vector databases hybrid search
Pinecone Weaviate Qdrant Milvus OpenSearch Elasticsearch FAISS
evaluation frameworks NDCG MRR MAP offline-to-online A/B test
LoRA QLoRA PEFT learning-to-rank XGBoost neural NLP IR recommendation system
"""

# Explicit, named JD-relevant skills. This is what lets the reasoning column
# say "matched on FAISS, hybrid search" instead of "strong tech score (0.34)".
# Adjust this list if the real job_description.md emphasizes different tools.
JD_KEY_SKILLS = {
    "embeddings", "retrieval", "ranking", "llm", "llms", "fine-tuning",
    "sentence-transformers", "bge", "e5", "vector database", "vector databases",
    "hybrid search", "pinecone", "weaviate", "qdrant", "milvus", "opensearch",
    "elasticsearch", "faiss", "ndcg", "mrr", "map", "a/b testing", "lora",
    "qlora", "peft", "learning-to-rank", "xgboost", "nlp", "information retrieval",
    "recommendation systems", "recommender systems", "pytorch", "transformers",
    "bert", "semantic search",
}

SERVICES_FIRMS = {"TCS", "Infosys", "Wipro", "Accenture", "Cognizant", "Capgemini", "Tech Mahindra"}

# Expanded vs v1 -- v1's list let "Operations Manager" reach rank 96.
IRRELEVANT_ROLES = [
    "marketing", " hr ", "human resources", "sales", "accountant", "accounting",
    "content writer", "content marketing", "graphic design", "civil engineer",
    "mechanical engineer", "customer support", "customer success", "operations",
    "office manager", "procurement", "logistics", "supply chain", "legal counsel",
    "recruiter", "talent acquisition", "administrative", "executive assistant",
    "finance manager", "internal audit", "receptionist", "front desk", "store manager",
]

AI_CORE_ROLES = [
    "ai", "machine learning", " ml ", "ml engineer", "backend", "data scientist",
    "data engineer", "search", "ranking", "retrieval", "nlp", "recommendation",
    "deep learning", "computer vision", "research scientist", "applied scientist",
    "mlops", "llm",
]

# Reference-date anchoring for inactivity math. Derived from the dataset
# itself by default (see compute_reference_date) so results don't silently
# shift if this script runs on a different wall-clock date during Stage 3.
INACTIVITY_PENALTY_START_DAYS = 60    # no penalty before this many days inactive
INACTIVITY_FULL_PENALTY_DAYS = 365    # penalty floors out beyond this
INACTIVITY_FLOOR_MULTIPLIER = 0.15    # never zero -- an exceptional match can still surface

STUFFING_SKILL_COUNT_THRESHOLD = 15
STUFFING_DEEP_SKILL_FRACTION = 0.30   # below this fraction of "real" (>=6mo) skills -> suspect
STUFFING_PENALTY_MULTIPLIER = 0.70

HONEYPOT_SCORE_THRESHOLD = 2          # corroborated-evidence threshold to exclude


# ---------------------------------------------------------------------------
# Loading
# ---------------------------------------------------------------------------

def open_maybe_gzip(path):
    """Detect gzip by magic bytes, not filename -- works regardless of how
    the file is named at Stage-3 reproduction time."""
    with open(path, "rb") as fh:
        magic = fh.read(2)
    if magic == b"\x1f\x8b":
        return gzip.open(path, "rt", encoding="utf-8")
    return open(path, "r", encoding="utf-8")


def find_candidates_file(explicit_path):
    if explicit_path:
        p = Path(explicit_path)
        if not p.exists():
            raise FileNotFoundError(f"--candidates path does not exist: {p}")
        return p
    for name in ("candidates.jsonl.gz", "candidates.jsonl", "candidates"):
        p = SCRIPT_DIR / name
        if p.exists():
            return p
    raise FileNotFoundError(
        "Could not find candidates.jsonl(.gz) next to this script. "
        "Pass --candidates /path/to/candidates.jsonl.gz explicitly."
    )


def read_docx_text(path):
    """Extract plain text from a .docx. Several hackathon bundles ship
    job_description.docx instead of .md -- v2 originally only looked for
    .md/.txt, which silently triggered the buzzword fallback even though
    the real JD was sitting right there as a .docx."""
    try:
        import docx  # python-docx
    except ImportError:
        print(
            "[WARN] Found a .docx job description but python-docx isn't "
            "installed, so it can't be read. Run: pip install python-docx",
            file=sys.stderr,
        )
        return None
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_job_description(explicit_path):
    """Read the real JD. Only fall back to the buzzword blob if no readable
    file is found at all -- v1 never even tried to read this file."""
    search_paths = [explicit_path] if explicit_path else []
    search_paths += [
        SCRIPT_DIR / "job_description.md",
        SCRIPT_DIR / "job_description.txt",
        SCRIPT_DIR / "job_description.docx",
    ]
    for p in search_paths:
        if not p:
            continue
        p = Path(p)
        if not p.exists():
            continue
        if p.suffix.lower() == ".docx":
            text = read_docx_text(p)
            if text:
                print(f"[ok] Loaded job description from {p} ({len(text)} chars).")
                return text
            continue  # couldn't read this one, keep looking
        text = p.read_text(encoding="utf-8")
        print(f"[ok] Loaded job description from {p} ({len(text)} chars).")
        return text
    print(
        "[WARN] No readable job_description.(md/txt/docx) found -- falling "
        "back to a curated keyword blob. Scores will NOT reflect the real "
        "JD. Pass --jd /path/to/job_description.<ext> to fix this.",
        file=sys.stderr,
    )
    return FALLBACK_JD_TEXT


def load_candidates(path):
    candidates = []
    with open_maybe_gzip(path) as f:
        for line in f:
            if not line.strip():
                continue
            candidates.append(json.loads(line))
    return candidates


def compute_reference_date(candidates, override):
    """Anchor 'today' to the freshest last_active_date in the dataset, so
    inactivity penalties are deterministic regardless of when/where this
    script actually runs. v1 hardcoded a constant date instead."""
    if override:
        return datetime.strptime(override, "%Y-%m-%d")
    max_date = None
    for cand in candidates:
        s = cand.get("redrob_signals", {}).get("last_active_date")
        if not s:
            continue
        try:
            d = datetime.strptime(s, "%Y-%m-%d")
        except ValueError:
            continue
        if max_date is None or d > max_date:
            max_date = d
    if max_date is None:
        print("[WARN] No last_active_date found anywhere; defaulting reference date to today.", file=sys.stderr)
        return datetime.combine(date.today(), datetime.min.time())
    print(f"[ok] Reference date anchored to freshest signal in data: {max_date.date()}")
    return max_date


# ---------------------------------------------------------------------------
# Per-candidate feature extraction
# ---------------------------------------------------------------------------

def detect_honeypot(profile, skills, signals):
    """Corroborated-evidence honeypot scoring, not a single brittle rule.
    Returns (score, reasons) -- candidate is excluded only if score crosses
    HONEYPOT_SCORE_THRESHOLD."""
    score = 0
    reasons = []

    years_exp = profile.get("years_of_experience", 0) or 0
    total_exp_months = years_exp * 12

    # Pattern 1: a single skill's duration exceeds total claimed experience.
    # Kept as a strong (+2) signal on its own since it's fairly unambiguous.
    if any((s.get("duration_months", 0) or 0) > total_exp_months for s in skills):
        score += 2
        reasons.append("skill duration exceeds total years of experience")

    # Pattern 2: "expert proficiency in N skills with 0 years used" -- the
    # exact example given in redrob_signals_doc. Cross-references the
    # skill_assessment_scores dict against the skills list's duration_months.
    assessment = signals.get("skill_assessment_scores", {}) or {}
    zero_duration_experts = 0
    for s in skills:
        name = s.get("name", "")
        dur = s.get("duration_months", 0) or 0
        assessed = assessment.get(name)
        if assessed is not None and assessed >= 85 and dur == 0:
            zero_duration_experts += 1
    if zero_duration_experts >= 3:
        score += 2
        reasons.append(f"{zero_duration_experts} skills assessed as expert-level with 0 months used")

    # Pattern 3 (guarded, schema-unconfirmed): "years of experience at a
    # company founded N years ago." candidate_schema.json wasn't available
    # at audit time, so this is a documented no-op unless the field exists --
    # adjust the key below once you confirm it (likely company_founded_year
    # or founded_year inside each career_history entry).
    # for h in career_history:
    #     founded = h.get("company_founded_year")
    #     start_year = h.get("start_date", "")[:4] if h.get("start_date") else None
    #     if founded and start_year and int(start_year) < int(founded):
    #         score += 2
    #         reasons.append(f"employment at {h.get('company')} predates its founding year")

    return score, reasons


def classify_role(title_lower):
    is_irrelevant = any(role in title_lower for role in IRRELEVANT_ROLES)
    is_core = any(role in title_lower for role in AI_CORE_ROLES)
    return is_irrelevant, is_core


def compute_matched_skills(skills, assessment):
    """Explicit, named overlap with JD_KEY_SKILLS, ranked by assessed score
    where available. This is what makes the reasoning column specific
    instead of generic."""
    matched = []
    for s in skills:
        name = s.get("name", "")
        if name.lower() in JD_KEY_SKILLS:
            matched.append((name, assessment.get(name, 0)))
    matched.sort(key=lambda x: x[1], reverse=True)
    return [m[0] for m in matched]


def inactivity_multiplier(days_inactive):
    if days_inactive <= INACTIVITY_PENALTY_START_DAYS:
        return 1.0
    if days_inactive >= INACTIVITY_FULL_PENALTY_DAYS:
        return INACTIVITY_FLOOR_MULTIPLIER
    frac = (days_inactive - INACTIVITY_PENALTY_START_DAYS) / (
        INACTIVITY_FULL_PENALTY_DAYS - INACTIVITY_PENALTY_START_DAYS
    )
    return 1.0 - (1.0 - INACTIVITY_FLOOR_MULTIPLIER) * frac


def compute_behavioral_score(signals):
    """Uses 7 of the 23 signals (v1 used 2), with open_to_work_flag as a
    soft multiplier rather than ignored entirely."""
    response_rate = max(0.0, min(1.0, signals.get("recruiter_response_rate", 0.0) or 0.0))
    github_raw = signals.get("github_activity_score", -1)
    github = max(0.0, (github_raw or 0)) / 100.0
    interview_completion = max(0.0, min(1.0, signals.get("interview_completion_rate", 0.0) or 0.0))
    profile_completeness = max(0.0, min(1.0, (signals.get("profile_completeness_score", 0) or 0) / 100.0))

    search_appearance = signals.get("search_appearance_30d", 0) or 0
    saved_by_recruiters = signals.get("saved_by_recruiters_30d", 0) or 0
    visibility = math.log1p(search_appearance + saved_by_recruiters)
    visibility_norm = min(visibility / 6.0, 1.0)  # empirical cap; revisit against real data distribution

    score = (
        response_rate * 0.30
        + github * 0.15
        + interview_completion * 0.15
        + profile_completeness * 0.10
        + visibility_norm * 0.15
    )

    open_to_work = signals.get("open_to_work_flag", False)
    score *= 1.0 if open_to_work else 0.70  # soft penalty, not a hard exclude

    notice_period = signals.get("notice_period_days", 0) or 0
    if notice_period > 60:
        score *= max(0.85, 1.0 - (notice_period - 60) / 400.0)

    return round(score, 6), response_rate, notice_period, open_to_work


def detect_stuffing(skills):
    skill_count = len(skills)
    if skill_count < STUFFING_SKILL_COUNT_THRESHOLD:
        return False, skill_count, 0
    deep_count = sum(1 for s in skills if (s.get("duration_months", 0) or 0) >= 6)
    frac = deep_count / skill_count
    return frac < STUFFING_DEEP_SKILL_FRACTION, skill_count, deep_count


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def process_candidates(raw_candidates, reference_date, debug_honeypots=0):
    rows = []
    honeypot_excluded = 0
    stuffing_flagged = 0
    debug_shown = 0

    for cand in raw_candidates:
        profile = cand.get("profile", {})
        history = cand.get("career_history", [])
        skills = cand.get("skills", [])
        signals = cand.get("redrob_signals", {})

        honeypot_score, honeypot_reasons = detect_honeypot(profile, skills, signals)
        if honeypot_score >= HONEYPOT_SCORE_THRESHOLD:
            honeypot_excluded += 1
            if debug_shown < debug_honeypots:
                debug_shown += 1
                years_exp = profile.get("years_of_experience", 0) or 0
                print(
                    f"[debug-honeypot] {cand.get('candidate_id')} | "
                    f"years_of_experience={years_exp} | reasons={honeypot_reasons} | "
                    f"skills={[(s.get('name'), s.get('duration_months')) for s in skills]}",
                    file=sys.stderr,
                )
            continue

        title = str(profile.get("current_title", "Unknown"))
        title_lower = title.lower()
        is_irrelevant_role, is_core_role = classify_role(title_lower)

        companies = sorted({h.get("company", "") for h in history if h.get("company")})
        only_services = bool(companies) and set(companies).issubset(SERVICES_FIRMS)

        skill_text = " ".join(s.get("name", "") for s in skills)
        history_text = " ".join(h.get("description", "") for h in history)
        agg_text = f"{profile.get('summary', '')} {history_text} {skill_text}"

        assessment = signals.get("skill_assessment_scores", {}) or {}
        matched_skills = compute_matched_skills(skills, assessment)

        is_stuffing_suspect, skill_count, deep_skill_count = detect_stuffing(skills)
        if is_stuffing_suspect:
            stuffing_flagged += 1

        last_active = signals.get("last_active_date")
        days_inactive = None
        if last_active:
            try:
                days_inactive = (reference_date - datetime.strptime(last_active, "%Y-%m-%d")).days
            except ValueError:
                days_inactive = None
        inactivity_mult = inactivity_multiplier(days_inactive) if days_inactive is not None else 1.0

        behavioral_score, response_rate, notice_period, open_to_work = compute_behavioral_score(signals)

        rows.append({
            "candidate_id": cand.get("candidate_id"),
            "title": title,
            "years_exp": profile.get("years_of_experience", 0) or 0,
            "agg_text": agg_text,
            "behavioral_score": behavioral_score,
            "response_rate": response_rate,
            "notice_period": notice_period,
            "open_to_work": open_to_work,
            "is_irrelevant_role": is_irrelevant_role,
            "is_core_role": is_core_role,
            "only_services": only_services,
            "companies": companies,
            "matched_skills": matched_skills,
            "is_stuffing_suspect": is_stuffing_suspect,
            "days_inactive": days_inactive,
            "inactivity_mult": inactivity_mult,
        })

    rate = honeypot_excluded / len(raw_candidates) if raw_candidates else 0
    print(f"[ok] {len(rows)} candidates retained; {honeypot_excluded} excluded as corroborated honeypots "
          f"({rate*100:.1f}% of pool); {stuffing_flagged} flagged as skill-stuffing suspects (penalized, not excluded).")
    if rate > 0.02:  # redrob_signals_doc.docx states ~80 honeypots in a 100k pool, i.e. ~0.08%
        print(
            f"[WARN] Honeypot exclusion rate ({rate*100:.1f}%) is far above the documented ~80-in-100k "
            "baseline (~0.08%). The detection rule is very likely over-firing on legitimate candidates, "
            "not finding real honeypots. Re-run with --debug-honeypots 10 and inspect the printed examples "
            "before trusting this output.",
            file=sys.stderr,
        )
    return pd.DataFrame(rows)


def rank_candidates(df, jd_text):
    vectorizer = TfidfVectorizer(
        stop_words="english",
        ngram_range=(1, 2),
        max_features=8000,
        sublinear_tf=True,  # dampens raw term-frequency stuffing -- the dataset's own stated trap
    )
    corpus = [jd_text] + df["agg_text"].tolist()
    tfidf_matrix = vectorizer.fit_transform(corpus)
    df["text_score"] = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

    df["skill_overlap_score"] = df["matched_skills"].apply(lambda m: min(len(m) / 5.0, 1.0))

    # Blend explicit skill overlap with general text similarity so a
    # candidate isn't purely at the mercy of TF-IDF's keyword sensitivity.
    df["content_score"] = (df["text_score"] * 0.6) + (df["skill_overlap_score"] * 0.4)
    if df["content_score"].max() > 0:
        df["content_score"] = df["content_score"] / df["content_score"].max()

    df.loc[df["is_irrelevant_role"], "content_score"] *= 0.1
    df.loc[df["only_services"], "content_score"] *= 0.7
    df.loc[df["is_core_role"], "content_score"] *= 1.2
    df.loc[df["is_stuffing_suspect"], "content_score"] *= STUFFING_PENALTY_MULTIPLIER
    df["content_score"] *= df["inactivity_mult"]

    if df["content_score"].max() > 0:
        df["content_score"] = df["content_score"] / df["content_score"].max()

    df["final_score"] = (df["content_score"] * 0.70) + (df["behavioral_score"] * 0.30)
    if df["final_score"].max() > 0:
        df["final_score"] = df["final_score"] / df["final_score"].max()

    top_100 = df.sort_values(by=["final_score", "candidate_id"], ascending=[False, True]).head(100).copy()
    top_100["rank"] = range(1, 101)
    return top_100


def generate_reasoning(row, total_rows):
    parts = [f"{row['title']} with {row['years_exp']:.1f} yrs experience."]

    if row["matched_skills"]:
        parts.append(f"Direct overlap with JD-relevant tooling: {', '.join(row['matched_skills'][:3])}.")
    else:
        parts.append(
            "No named skill in their profile directly matches the JD's listed tools; "
            "ranked on broader text similarity and behavioral signals only."
        )

    if row["response_rate"] >= 0.6:
        parts.append(f"Recruiter response rate is strong at {row['response_rate']*100:.0f}%.")
    elif row["response_rate"] <= 0.25:
        parts.append(f"Recruiter response rate is low ({row['response_rate']*100:.0f}%), a real availability concern.")

    if row["days_inactive"] is not None and row["days_inactive"] > 90:
        parts.append(f"Last active on the platform {row['days_inactive']} days ago.")

    if row["only_services"]:
        co = ", ".join(row["companies"][:2])
        parts.append(f"Career history is limited to services firms ({co}), which may mean less direct ownership of production systems.")

    if not row["is_core_role"]:
        if row["is_irrelevant_role"]:
            parts.append("Title reads as unrelated to the JD's core AI/ML scope -- retained only on a behavioral-signal edge case; worth a manual check.")
        else:
            parts.append("Title is adjacent rather than core AI/ML, so fit is partial.")

    if not row["open_to_work"]:
        parts.append("Not currently flagged open-to-work, which may slow outreach.")

    if row["notice_period"] and row["notice_period"] > 60:
        parts.append(f"Notice period is {row['notice_period']} days.")

    if row["is_stuffing_suspect"]:
        parts.append("Skill list is unusually long relative to assessed depth -- treated with caution.")

    if row["rank"] > total_rows * 0.7 and len(parts) <= 2:
        parts.append("Included near the bottom of the top 100 primarily on aggregate score; no single standout strength.")

    return " ".join(parts)


def self_check(submission_df):
    """Mirrors the official validator's core rules so a logic bug fails
    loudly here, not silently at Stage 1."""
    errors = []
    if len(submission_df) != 100:
        errors.append(f"Expected 100 rows, got {len(submission_df)}")
    if sorted(submission_df["rank"].tolist()) != list(range(1, 101)):
        errors.append("Ranks are not exactly 1..100 each-once")
    if submission_df["candidate_id"].duplicated().any():
        errors.append("Duplicate candidate_id found")
    scores_by_rank = submission_df.sort_values("rank")["score"].tolist()
    if any(scores_by_rank[i] < scores_by_rank[i + 1] for i in range(len(scores_by_rank) - 1)):
        errors.append("Score is not non-increasing by rank")
    if errors:
        raise AssertionError("Internal self-check failed:\n  " + "\n  ".join(errors))
    print("[ok] Internal self-check passed (mirrors validate_submission.py core rules).")


def main():
    parser = argparse.ArgumentParser(description="Redrob Hackathon candidate ranker")
    parser.add_argument("--candidates", help="Path to candidates.jsonl or candidates.jsonl.gz")
    parser.add_argument("--jd", help="Path to job_description.md")
    parser.add_argument("--out", help="Output CSV path", default=str(SCRIPT_DIR / "team_devwithsk.csv"))
    parser.add_argument("--reference-date", help="Override reference date, format YYYY-MM-DD")
    parser.add_argument(
        "--debug-honeypots", type=int, default=0,
        help="Print this many excluded honeypot examples (candidate_id, years_of_experience, "
             "reasons, raw skill durations) to help diagnose an unexpectedly high exclusion rate.",
    )
    args = parser.parse_args()

    t0 = time.time()

    candidates_path = find_candidates_file(args.candidates)
    print(f"[ok] Loading candidates from {candidates_path} ...")
    raw_candidates = load_candidates(candidates_path)
    print(f"[ok] Loaded {len(raw_candidates)} candidates.")

    jd_text = load_job_description(args.jd)
    reference_date = compute_reference_date(raw_candidates, args.reference_date)

    df = process_candidates(raw_candidates, reference_date, debug_honeypots=args.debug_honeypots)
    df_ranked = rank_candidates(df, jd_text)
    df_ranked["reasoning"] = df_ranked.apply(lambda r: generate_reasoning(r, len(df_ranked)), axis=1)

    submission_df = df_ranked[["candidate_id", "rank", "final_score", "reasoning"]].rename(
        columns={"final_score": "score"}
    )

    self_check(submission_df)

    submission_df.to_csv(args.out, index=False, encoding="utf-8")
    elapsed = time.time() - t0
    print(f"[ok] Wrote {len(submission_df)} rows to {args.out}")
    print(f"[ok] Total runtime: {elapsed:.1f}s (must stay under 300s / 5min per Section 3)")


if __name__ == "__main__":
    main()